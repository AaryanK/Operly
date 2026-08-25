"""Deterministic artifact authoring and format conversion capabilities.

These primitives are deliberately separate from ``files.process``. The attachment
processor is for understanding untrusted input; this provider turns already-known
structured/text content into exact requested output artifacts without asking a model
to analyze a synthetic attachment first.
"""
from __future__ import annotations

import csv
import io
import re
import tempfile
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

from packages.artifacts.service import ArtifactService, artifact_json, artifact_scope_from_context
from packages.capabilities.contracts import ApprovalPolicy, CapabilityDefinition, CapabilityResult
from packages.capabilities.providers import BaseProvider


_DOCUMENT_FORMATS = {"pdf", "docx", "txt", "markdown"}
_SPREADSHEET_FORMATS = {"xlsx", "csv"}
_CONVERT_FORMATS = _DOCUMENT_FORMATS | _SPREADSHEET_FORMATS


def _run_id(context) -> str | None:
    invocation = context.invocation if isinstance(context.invocation, dict) else {}
    metadata = invocation.get("metadata") if isinstance(invocation.get("metadata"), dict) else {}
    return str(metadata.get("runtime_run_id") or "").strip() or None


def _safe_filename(value: str | None, *, stem_default: str, extension: str) -> str:
    raw = str(value or "").strip()
    stem = Path(raw).stem if raw else stem_default
    stem = re.sub(r"[^A-Za-z0-9._-]+", "-", stem).strip("-.")[:120] or stem_default
    return f"{stem}.{extension}"


def _safe_cell(value: Any) -> Any:
    if value is None or isinstance(value, (bool, int, float)):
        return value
    text = str(value)
    return "'" + text if text.startswith(("=", "+", "-", "@")) else text


def _document_content_type(output_format: str) -> str:
    return {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "txt": "text/plain; charset=utf-8",
        "markdown": "text/markdown; charset=utf-8",
    }[output_format]


def _spreadsheet_content_type(output_format: str) -> str:
    return {
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "csv": "text/csv; charset=utf-8",
    }[output_format]


def _write_document(path: Path, *, output_format: str, title: str, content: str) -> None:
    if output_format == "pdf":
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError as exc:
            raise RuntimeError("PDF output dependency is unavailable") from exc
        styles = getSampleStyleSheet()
        story = []
        if title:
            story.extend([Paragraph(escape(title), styles["Title"]), Spacer(1, 12)])
        for line in content.splitlines() or [""]:
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 7))
                continue
            style = styles["BodyText"]
            text = stripped
            if stripped.startswith("### "):
                style, text = styles["Heading3"], stripped[4:]
            elif stripped.startswith("## "):
                style, text = styles["Heading2"], stripped[3:]
            elif stripped.startswith("# "):
                style, text = styles["Heading1"], stripped[2:]
            elif stripped.startswith(("- ", "* ")):
                text = "• " + stripped[2:]
            story.append(Paragraph(escape(text), style))
        SimpleDocTemplate(str(path), pagesize=letter).build(story)
        return

    if output_format == "docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX output dependency is unavailable") from exc
        doc = Document()
        if title:
            doc.add_heading(title, 0)
        for line in content.splitlines() or [""]:
            stripped = line.strip()
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith(("- ", "* ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)
        doc.save(path)
        return

    path.write_text(content, encoding="utf-8")


def _write_spreadsheet(
    path: Path,
    *,
    output_format: str,
    columns: list[str],
    rows: list[dict[str, Any]],
    title: str,
) -> None:
    if output_format == "xlsx":
        try:
            from openpyxl import Workbook
        except ImportError as exc:
            raise RuntimeError("XLSX output dependency is unavailable") from exc
        wb = Workbook()
        ws = wb.active
        ws.title = (title or "OPERLY Results")[:31]
        ws.append(columns)
        for row in rows:
            ws.append([_safe_cell(row.get(column)) for column in columns])
        wb.save(path)
        return

    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(columns)
        for row in rows:
            writer.writerow([_safe_cell(row.get(column)) for column in columns])


def _decode_text(raw: bytes) -> str:
    try:
        return raw.decode("utf-8-sig")
    except UnicodeDecodeError as error:
        raise ValueError("source artifact is not valid UTF-8 text") from error


def _parse_table(raw: bytes, *, suffix: str, content_type: str) -> tuple[list[str], list[dict[str, Any]]]:
    text = _decode_text(raw)
    delimiter = "\t" if suffix == ".tsv" or content_type == "text/tab-separated-values" else ","
    material = list(csv.reader(io.StringIO(text), delimiter=delimiter))
    if not material:
        raise ValueError("source table is empty")
    columns = [str(item) for item in material[0]]
    if not columns or any(not column.strip() for column in columns):
        raise ValueError("source table must have non-empty header columns")
    rows = [
        {column: (record[index] if index < len(record) else None) for index, column in enumerate(columns)}
        for record in material[1:]
    ]
    return columns, rows


class FileAuthoringProvider(BaseProvider):
    """Create exact user-facing artifacts from known content without perception calls."""

    name = "operly_file_authoring"
    capabilities = (
        CapabilityDefinition(
            "files.create_document",
            "files_create_document",
            (
                "Create a durable PDF, DOCX, TXT or Markdown document directly from content already known to the agent. "
                "Use this for executive summaries and authored reports. Do not create a fake attachment or base64 wrapper."
            ),
            {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "maxLength": 200},
                    "content": {"type": "string", "minLength": 1, "maxLength": 100000},
                    "output_format": {"type": "string", "enum": sorted(_DOCUMENT_FORMATS)},
                    "filename": {"type": "string", "maxLength": 255},
                },
                "required": ["content", "output_format"],
                "additionalProperties": False,
            },
            {"type": "object"},
            risk_level="low",
            permissions=("files:process",),
            approval_policy=ApprovalPolicy.AUTO,
            execution_timeout_seconds=120,
            reversible=True,
            category="files",
            display_name="Create document",
            tags=frozenset({"files", "documents", "pdf", "docx", "reports", "artifacts"}),
            semantic_operations=frozenset(
                {
                    "create pdf from content",
                    "create executive summary pdf",
                    "write report to pdf",
                    "create docx from content",
                    "author document",
                }
            ),
        ),
        CapabilityDefinition(
            "files.create_spreadsheet",
            "files_create_spreadsheet",
            (
                "Create a durable XLSX or CSV directly from explicit structured rows. The declared columns are the actual workbook schema. "
                "Use this when the user asks for Excel/spreadsheet output from data already gathered by the agent."
            ),
            {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "maxLength": 200},
                    "columns": {
                        "type": "array",
                        "minItems": 1,
                        "maxItems": 50,
                        "items": {"type": "string", "minLength": 1, "maxLength": 120},
                    },
                    "rows": {"type": "array", "maxItems": 2000, "items": {"type": "object"}},
                    "output_format": {"type": "string", "enum": sorted(_SPREADSHEET_FORMATS)},
                    "filename": {"type": "string", "maxLength": 255},
                },
                "required": ["columns", "rows", "output_format"],
                "additionalProperties": False,
            },
            {"type": "object"},
            risk_level="low",
            permissions=("files:process",),
            approval_policy=ApprovalPolicy.AUTO,
            execution_timeout_seconds=120,
            reversible=True,
            category="files",
            display_name="Create spreadsheet",
            tags=frozenset({"files", "spreadsheet", "xlsx", "csv", "table", "artifacts"}),
            semantic_operations=frozenset(
                {
                    "create excel workbook",
                    "create xlsx spreadsheet",
                    "create spreadsheet from rows",
                    "export structured data to excel",
                }
            ),
        ),
        CapabilityDefinition(
            "files.convert",
            "files_convert",
            (
                "Deterministically convert one durable artifact without asking a model to understand its contents. "
                "Use for format-only operations such as image to PDF, text to PDF/DOCX, or CSV/TSV to XLSX."
            ),
            {
                "type": "object",
                "properties": {
                    "artifact_id": {"type": "string", "minLength": 1, "maxLength": 36},
                    "output_format": {"type": "string", "enum": sorted(_CONVERT_FORMATS)},
                    "filename": {"type": "string", "maxLength": 255},
                    "title": {"type": "string", "maxLength": 200},
                },
                "required": ["artifact_id", "output_format"],
                "additionalProperties": False,
            },
            {"type": "object"},
            risk_level="low",
            permissions=("files:process",),
            approval_policy=ApprovalPolicy.AUTO,
            execution_timeout_seconds=120,
            reversible=True,
            category="files",
            display_name="Convert file",
            tags=frozenset({"files", "convert", "pdf", "xlsx", "artifacts", "deterministic"}),
            semantic_operations=frozenset(
                {
                    "convert image to pdf",
                    "convert file format",
                    "convert csv to xlsx",
                    "convert text to pdf",
                }
            ),
        ),
    )

    async def _persist(
        self,
        context,
        *,
        path: Path,
        filename: str,
        content_type: str,
        generated_by: str,
        metadata: dict[str, Any],
        parent_artifact_id: str | None = None,
    ) -> dict[str, Any]:
        row = await ArtifactService(context.db).create_path(
            artifact_scope_from_context(context),
            path,
            filename=filename,
            content_type=content_type,
            source="agent_generated",
            created_by=context.actor_id,
            run_id=_run_id(context),
            parent_artifact_id=parent_artifact_id,
            metadata={"generated_by": generated_by, **metadata},
        )
        return artifact_json(row)

    async def _create_document(self, context, arguments: dict[str, Any]) -> CapabilityResult:
        output_format = str(arguments.get("output_format") or "").lower()
        if output_format not in _DOCUMENT_FORMATS:
            raise ValueError("unsupported document output format")
        title = str(arguments.get("title") or "")[:200]
        content = str(arguments.get("content") or "")[:100000]
        if not content.strip():
            raise ValueError("document content is required")
        extension = "md" if output_format == "markdown" else output_format
        filename = _safe_filename(arguments.get("filename"), stem_default="operly-report", extension=extension)
        with tempfile.TemporaryDirectory(prefix="operly-author-document-") as temp_dir:
            path = Path(temp_dir) / filename
            _write_document(path, output_format=output_format, title=title, content=content)
            artifact = await self._persist(
                context,
                path=path,
                filename=filename,
                content_type=_document_content_type(output_format),
                generated_by="files.create_document",
                metadata={"artifact_kind": "document", "output_format": output_format, "title": title},
            )
        return CapabilityResult(
            True,
            True,
            {
                "artifact_id": artifact["artifact_id"],
                "artifact_ids": [artifact["artifact_id"]],
                "artifacts": [artifact],
                "artifact_kind": "document",
                "output_format": output_format,
                "title": title,
                "transport": "artifact_authored_v1",
                "side_effects": False,
            },
            artifact["artifact_id"],
        )

    async def _create_spreadsheet(self, context, arguments: dict[str, Any]) -> CapabilityResult:
        output_format = str(arguments.get("output_format") or "").lower()
        if output_format not in _SPREADSHEET_FORMATS:
            raise ValueError("unsupported spreadsheet output format")
        columns = [str(item).strip()[:120] for item in arguments.get("columns") or []]
        if not columns or any(not item for item in columns):
            raise ValueError("spreadsheet columns are required")
        if len(set(columns)) != len(columns):
            raise ValueError("spreadsheet columns must be unique")
        rows = list(arguments.get("rows") or [])
        if any(not isinstance(row, dict) for row in rows):
            raise ValueError("spreadsheet rows must be objects")
        if len(rows) > 2000:
            raise ValueError("spreadsheet supports at most 2000 rows")
        title = str(arguments.get("title") or "OPERLY Results")[:200]
        filename = _safe_filename(arguments.get("filename"), stem_default="operly-results", extension=output_format)
        with tempfile.TemporaryDirectory(prefix="operly-author-sheet-") as temp_dir:
            path = Path(temp_dir) / filename
            _write_spreadsheet(path, output_format=output_format, columns=columns, rows=rows, title=title)
            artifact = await self._persist(
                context,
                path=path,
                filename=filename,
                content_type=_spreadsheet_content_type(output_format),
                generated_by="files.create_spreadsheet",
                metadata={
                    "artifact_kind": "spreadsheet",
                    "output_format": output_format,
                    "columns": columns,
                    "row_count": len(rows),
                },
            )
        return CapabilityResult(
            True,
            True,
            {
                "artifact_id": artifact["artifact_id"],
                "artifact_ids": [artifact["artifact_id"]],
                "artifacts": [artifact],
                "artifact_kind": "spreadsheet",
                "output_format": output_format,
                "columns": columns,
                "row_count": len(rows),
                "transport": "artifact_authored_v1",
                "side_effects": False,
            },
            artifact["artifact_id"],
        )

    async def _convert(self, context, arguments: dict[str, Any]) -> CapabilityResult:
        artifact_id = str(arguments.get("artifact_id") or "").strip()
        output_format = str(arguments.get("output_format") or "").lower()
        if output_format not in _CONVERT_FORMATS:
            raise ValueError("unsupported conversion output format")
        service = ArtifactService(context.db)
        scope = artifact_scope_from_context(context)
        source = await service.get(scope, artifact_id)
        raw = await service.read_bytes(scope, source.id)
        source_type = str(source.content_type or "").split(";", 1)[0].strip().lower()
        suffix = Path(source.filename).suffix.lower()
        extension = "md" if output_format == "markdown" else output_format
        filename = _safe_filename(
            arguments.get("filename"), stem_default=Path(source.filename).stem or "operly-converted", extension=extension
        )
        title = str(arguments.get("title") or Path(source.filename).stem)[:200]
        conversion: dict[str, Any]
        with tempfile.TemporaryDirectory(prefix="operly-convert-") as temp_dir:
            path = Path(temp_dir) / filename
            if output_format == "pdf" and (
                source_type.startswith("image/")
                or suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
            ):
                try:
                    from PIL import Image
                except ImportError as exc:
                    raise RuntimeError("Image conversion dependency is unavailable") from exc
                with Image.open(io.BytesIO(raw)) as image:
                    image.convert("RGB").save(path, "PDF", resolution=144.0)
                conversion = {"conversion": "image_to_pdf", "source_kind": "image"}
            elif suffix in {".txt", ".md", ".markdown", ".log"} or source_type in {"text/plain", "text/markdown"}:
                text = _decode_text(raw)
                if output_format not in _DOCUMENT_FORMATS:
                    raise ValueError(f"unsupported deterministic conversion from text to {output_format}")
                _write_document(path, output_format=output_format, title=title, content=text)
                conversion = {"conversion": f"text_to_{output_format}", "source_kind": "text"}
            elif suffix in {".csv", ".tsv"} or source_type in {"text/csv", "text/tab-separated-values"}:
                columns, rows = _parse_table(raw, suffix=suffix, content_type=source_type)
                if output_format in _SPREADSHEET_FORMATS:
                    _write_spreadsheet(path, output_format=output_format, columns=columns, rows=rows, title=title)
                elif output_format in _DOCUMENT_FORMATS:
                    lines = [" | ".join(columns)] + [
                        " | ".join(str(row.get(column) or "") for column in columns) for row in rows
                    ]
                    _write_document(path, output_format=output_format, title=title, content="\n".join(lines))
                else:
                    raise ValueError(f"unsupported deterministic conversion from table to {output_format}")
                conversion = {
                    "conversion": f"table_to_{output_format}",
                    "source_kind": "table",
                    "columns": columns,
                    "row_count": len(rows),
                }
            else:
                raise ValueError(
                    f"unsupported deterministic conversion from {source_type or suffix or 'unknown'} to {output_format}"
                )

            content_type = (
                _document_content_type(output_format)
                if output_format in _DOCUMENT_FORMATS
                else _spreadsheet_content_type(output_format)
            )
            artifact = await self._persist(
                context,
                path=path,
                filename=filename,
                content_type=content_type,
                generated_by="files.convert",
                parent_artifact_id=source.id,
                metadata={
                    "artifact_kind": "conversion",
                    "output_format": output_format,
                    "source_artifact_id": source.id,
                    **conversion,
                },
            )
        return CapabilityResult(
            True,
            True,
            {
                "artifact_id": artifact["artifact_id"],
                "artifact_ids": [artifact["artifact_id"]],
                "artifacts": [artifact],
                "artifact_kind": "conversion",
                "input_artifact_id": source.id,
                "source_content_type": source.content_type,
                "output_format": output_format,
                **conversion,
                "transport": "artifact_conversion_v1",
                "side_effects": False,
            },
            artifact["artifact_id"],
        )

    async def execute(self, context, capability_name: str, arguments: dict[str, Any]) -> CapabilityResult:
        try:
            if capability_name == "files.create_document":
                return await self._create_document(context, arguments)
            if capability_name == "files.create_spreadsheet":
                return await self._create_spreadsheet(context, arguments)
            if capability_name == "files.convert":
                return await self._convert(context, arguments)
            return CapabilityResult(False, False, {"reason": "unsupported_capability"})
        except (ValueError, LookupError, RuntimeError) as error:
            return CapabilityResult(False, False, {"reason": "file_authoring_failed", "message": str(error)[:1000]})

    async def verify(self, context, capability_name: str, arguments: dict[str, Any], result: CapabilityResult) -> CapabilityResult:
        del arguments
        if capability_name not in {"files.create_document", "files.create_spreadsheet", "files.convert"}:
            return CapabilityResult(False, False, {"reason": "unsupported_capability"})
        if not result.success:
            return CapabilityResult(False, result.changed, result.evidence)
        artifact_ids = list(result.evidence.get("artifact_ids") or [])
        if len(artifact_ids) != 1:
            return CapabilityResult(False, result.changed, {"reason": "authored_artifact_missing"})
        try:
            rows = await ArtifactService(context.db).get_many(
                artifact_scope_from_context(context), artifact_ids, max_items=1
            )
        except (LookupError, ValueError):
            return CapabilityResult(False, result.changed, {"reason": "authored_artifact_missing"})
        if len(rows) != 1:
            return CapabilityResult(False, result.changed, {"reason": "authored_artifact_missing"})
        artifact = artifact_json(rows[0])
        evidence = {
            "artifact_id": artifact_ids[0],
            "artifact_ids": artifact_ids,
            "artifacts": [artifact],
            "artifact_kind": result.evidence.get("artifact_kind"),
            "output_format": result.evidence.get("output_format"),
            "columns": result.evidence.get("columns") or [],
            "row_count": result.evidence.get("row_count"),
            "input_artifact_id": result.evidence.get("input_artifact_id"),
            "conversion": result.evidence.get("conversion"),
            "transport": result.evidence.get("transport"),
            "artifacts_persisted": True,
            "side_effects": False,
        }
        return CapabilityResult(True, result.changed, evidence, result.external_reference)
